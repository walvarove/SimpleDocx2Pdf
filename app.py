from flask import Flask, request, jsonify, send_file
import os
import uuid
import json
import subprocess
import re
import zipfile
import logging
from bs4 import BeautifulSoup
from docx import Document

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = Flask(__name__)

# Configuration
app.config['UPLOAD_FOLDER'] = '/tmp/docx2pdf'
app.config['MAX_CONTENT_LENGTH'] = 10 * 1024 * 1024  # 10 MB
app.config['API_TOKEN'] = os.environ.get('API_TOKEN', 'Bkbxl2376APMv99y77HmGWk2teBODjXO')

# Ensure the upload folder exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

def docx_to_html(docx_file, html_file):
    """Convert DOCX to HTML using Pandoc with metadata"""
    try:
        subprocess.run([
            "pandoc",
            docx_file,
            "-f", "docx",
            "-t", "html5",
            "--metadata", "title= ",  # Empty title to prevent auto-generation
            "--standalone",
            "--embed-resources",
            "--extract-media", ".",  # Extract media files if any
            "-o", html_file
        ], check=True, capture_output=True, text=True)

        # Read the generated HTML
        with open(html_file, 'r', encoding='utf-8') as f:
            content = f.read()

        # Remove or replace the auto-generated title
        content = re.sub(r'<title>.*?</title>', '<title></title>', content)

        # Write back the modified HTML
        with open(html_file, 'w', encoding='utf-8') as f:
            f.write(content)

    except subprocess.CalledProcessError as e:
        logger.error(f"Pandoc conversion error: {e}")
        logger.error(f"Command output: {e.stdout}\n{e.stderr}")
        raise
    except Exception as e:
        logger.error(f"Error in docx_to_html: {e}")
        raise

def html_to_docx(html_file, output_docx, reference_doc=None):
    """Convert HTML back to DOCX using Pandoc"""
    cmd = ["pandoc", "-f", "html", "-t", "docx"]
    if reference_doc:
        cmd.extend(["--reference-doc", reference_doc])
    cmd.extend([html_file, "-o", output_docx])
    subprocess.run(cmd, check=True)

def convert_to_pdf(input_path, output_path):
    """Convert DOCX to PDF using LibreOffice"""
    cmd = [
        'libreoffice',
        '--headless',
        '--convert-to', 'pdf',
        '--outdir', os.path.dirname(output_path),
        input_path
    ]
    process = subprocess.run(cmd, capture_output=True, text=True)
    
    if process.returncode == 0:
        base_name = os.path.basename(input_path)
        base_name_without_ext = os.path.splitext(base_name)[0]
        libreoffice_output = os.path.join(os.path.dirname(output_path), f"{base_name_without_ext}.pdf")
        
        if libreoffice_output != output_path:
            os.rename(libreoffice_output, output_path)
        return True
    return False

def ensure_list_style(doc):
    """Ensure the document has the required list style"""
    try:
        doc.styles['List Bullet']
    except KeyError:
        # If style doesn't exist, create it
        style = doc.styles.add_style('List Bullet', 1)  # 1 is WD_STYLE_TYPE.PARAGRAPH
        style.base_style = doc.styles['Normal']
        # Add bullet point
        fmt = style.paragraph_format
        fmt.left_indent = 914400  # 1 inch in EMUs
        fmt.first_line_indent = -457200  # -0.5 inch
        fmt.space_after = 0
        fmt.line_spacing = 1

def process_html_to_docx_content(doc, html_content):
    """Process HTML content and add it to the document"""
    ensure_list_style(doc)
    soup = BeautifulSoup(html_content, 'html.parser')
    
    if soup.find('ul'):
        for li in soup.find_all('li'):
            # Create paragraph with bullet
            paragraph = doc.add_paragraph()
            paragraph.style = 'List Bullet'
            paragraph.paragraph_format.left_indent = 1440  # 1 inch = 1440 twips
            run = paragraph.add_run('• ')  # Add bullet point manually
            
            # Process the content of each list item
            for content in li.children:
                if isinstance(content, str):
                    text = content.strip()
                    if text:
                        run = paragraph.add_run(text)
                else:
                    text = content.get_text().strip()
                    if text:
                        run = paragraph.add_run(text)
                        if content.name == 'b':
                            run.bold = True
                        if content.name == 'u':
                            run.underline = True
    else:
        # For non-list HTML content
        paragraph = doc.add_paragraph()
        paragraph.add_run(soup.get_text())
    return doc

def replace_variables(template_doc, replacements):
    """Replace variables in the document"""
    # First pass: collect all replacements
    replacements_to_process = []
    
    # Process all paragraphs including headers
    for i, paragraph in enumerate(template_doc.paragraphs):
        logger.info(f"Processing paragraph {i}: {paragraph.text}")
        for key, value in replacements.items():
            placeholder = '{' + key + '}'
            
            if placeholder in paragraph.text:
                logger.info(f"Found placeholder {placeholder} in paragraph {i}")
                if isinstance(value, str) and re.search(r'<[^>]+>', value):
                    # Store the paragraph index, placeholder and HTML content
                    replacements_to_process.append((i, placeholder, value))
                else:
                    # For regular text replacements
                    for run in paragraph.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, str(value))
    
    # Second pass: process HTML content in reverse order
    for idx, placeholder, html_content in sorted(replacements_to_process, key=lambda x: x[0], reverse=True):
        try:
            # Get the paragraph that contains the placeholder
            target_paragraph = template_doc.paragraphs[idx]
            
            # Create a new document for the HTML content
            temp_doc = Document()
            process_html_to_docx_content(temp_doc, html_content)
            
            # Get the parent element and index of the target paragraph
            parent = target_paragraph._element.getparent()
            insert_idx = parent.index(target_paragraph._element)
            
            # Insert each paragraph from the temp document at the correct position
            for html_para in temp_doc.paragraphs:
                # Create a new paragraph in the main document
                new_para = template_doc.add_paragraph()
                
                # Copy content and formatting
                for run in html_para.runs:
                    new_run = new_para.add_run(run.text)
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    new_run.underline = run.underline
                new_para.style = html_para.style
                
                # Set paragraph properties
                if html_para.style and html_para.style.name == 'List Bullet':
                    new_para.paragraph_format.left_indent = 1440  # 1 inch = 1440 twips
                
                # Insert at the correct position
                parent.insert(insert_idx, new_para._element)
                insert_idx += 1
            
            # Remove the original placeholder paragraph
            parent.remove(target_paragraph._element)
            
            logger.info(f"Successfully replaced placeholder at paragraph {idx} with HTML content")
        except Exception as e:
            logger.error(f"Error processing HTML replacement at paragraph {idx}: {str(e)}")
            raise

@app.route('/health', methods=['GET'])
def health_check():
    return jsonify(status='ok')

@app.route('/api/v1/parse_template', methods=['POST'])
def process_template_to_pdf():
    if request.headers.get('X-API-Token') != app.config['API_TOKEN']:
        return jsonify(error='Invalid or missing API token'), 401

    if 'template' not in request.files:
        return jsonify(error='No template file provided'), 400

    template_file = request.files['template']
    if not template_file.filename.endswith('.docx'):
        return jsonify(error='Invalid template file. Must be a .docx file'), 400

    unique_id = str(uuid.uuid4())
    template_path = os.path.join(app.config['UPLOAD_FOLDER'], f'{unique_id}_template.docx')
    output_docx = os.path.join(app.config['UPLOAD_FOLDER'], f'{unique_id}_output.docx')
    output_pdf = os.path.join(app.config['UPLOAD_FOLDER'], f'{unique_id}_output.pdf')
    zip_path = os.path.join(app.config['UPLOAD_FOLDER'], f'{unique_id}_output.zip')
    temp_files = [template_path, output_docx, output_pdf, zip_path]

    try:
        # Save template file
        template_file.save(template_path)
        
        # Load the template
        doc = Document(template_path)
        
        # Log template content for debugging
        logger.info("Template content:")
        for i, para in enumerate(doc.paragraphs):
            logger.info(f"Paragraph {i}: {para.text}")
        
        # Process variables
        data = json.loads(request.form.get('data', '{}'))
        logger.info(f"Processing data: {data}")
        
        # Process HTML content
        replace_variables(doc, data)
        
        # Log final content for debugging
        logger.info("Final document content:")
        for i, para in enumerate(doc.paragraphs):
            logger.info(f"Paragraph {i}: {para.text}")
        
        # Save the document
        doc.save(output_docx)
        
        # Convert to PDF
        if convert_to_pdf(output_docx, output_pdf):
            output_filename = request.form.get('filename', 'output.docx')
            if not output_filename.endswith('.docx'):
                output_filename += '.docx'
            
            with zipfile.ZipFile(zip_path, 'w', compression=zipfile.ZIP_DEFLATED) as zipf:
                zipf.write(output_docx, output_filename)
                pdf_filename = os.path.splitext(output_filename)[0] + '.pdf'
                zipf.write(output_pdf, pdf_filename)
            
            return send_file(zip_path, as_attachment=True, download_name=f'{os.path.splitext(output_filename)[0]}.zip')
        else:
            return jsonify(error="PDF conversion failed"), 500

    except Exception as e:
        logger.error(f"Error: {str(e)}")
        return jsonify(error=str(e)), 500
    
    finally:
        # Clean up temporary files
        for file_path in temp_files:
            if os.path.exists(file_path):
                try:
                    os.remove(file_path)
                except Exception as e:
                    logger.error(f"Error cleaning up {file_path}: {str(e)}")

@app.route('/api/v1/html_to_docx', methods=['POST'])
def process_html_to_docx():
    if request.headers.get('X-API-Token') != app.config['API_TOKEN']:
        return jsonify(error='Invalid or missing API token'), 401

    if 'html' not in request.form:
        return jsonify(error='No HTML content provided'), 400
    
    html_content = request.form['html']
    unique_id = str(uuid.uuid4())
    temp_html = os.path.join(app.config['UPLOAD_FOLDER'], f'{unique_id}_temp.html')
    output_docx = os.path.join(app.config['UPLOAD_FOLDER'], f'{unique_id}_output.docx')
    
    try:
        with open(temp_html, 'w', encoding='utf-8') as file:
            file.write(html_content)
        
        html_to_docx(temp_html, output_docx)
        
        return send_file(output_docx, as_attachment=True, download_name=f'{unique_id}_output.docx')
    
    except Exception as e:
        logger.error(f"Error: {str(e)}")
        return jsonify(error=str(e)), 500
    
    finally:
        if os.path.exists(temp_html):
            try:
                os.remove(temp_html)
            except Exception as e:
                logger.error(f"Error cleaning up {temp_html}: {str(e)}")

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 8080))
    app.run(host='0.0.0.0', port=port) 